import customtkinter as ctk
from tkinter import filedialog, messagebox
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from tkinterweb import HtmlFrame
import os
import re

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")


class KSeFApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("KSeF Podgląd i Opis Merytoryczny EZD PUW / EZD RP - System Samorządowy wersja beta 1.0.01.04.2026")
        self.geometry("1600x1000")
        self.invoice_data = {}
        self.invoice_items = []
        self.vat_summary = []
        self.podmioty3 = []
        self.ns = {}
        self.filename_ksef_id = ""
        self.setup_ui()

    def setup_ui(self):
        self.sidebar = ctk.CTkFrame(self, width=280, corner_radius=0)
        self.sidebar.pack(side="left", fill="y")

        self.logo = ctk.CTkLabel(self.sidebar, text="KSeFik SPK", font=ctk.CTkFont(size=22, weight="bold"))
        self.logo.pack(pady=30)

        self.btn_load = ctk.CTkButton(self.sidebar, text="Wczytaj Fakturę XML", command=self.load_xml, height=45)
        self.btn_load.pack(pady=10, padx=20)

        self.btn_export = ctk.CTkButton(self.sidebar, text="Generuj Opis Faktury",
                                        command=self.export_word, fg_color="#10b981", hover_color="#059669", height=45)
        self.btn_export.pack(pady=10, padx=20)

        self.status_label = ctk.CTkLabel(self.sidebar, text="Status: Gotowy", text_color="gray")
        self.status_label.pack(side="bottom", pady=20)

        self.container = ctk.CTkFrame(self, fg_color="#1a1a1a")
        self.container.pack(side="right", fill="both", expand=True)

        self.preview_frame = HtmlFrame(self.container, messages_enabled=False)
        self.preview_frame.pack(fill="both", expand=True, padx=20, pady=20)
        self.show_welcome_screen()

    def show_welcome_screen(self):
        self.preview_frame.load_html(
            "<body style='background:#1a1a1a; color:white; font-family:sans-serif; text-align:center; padding-top:250px;'>"
            "<h1>System Podglądu Faktur KSeF</h1>"
            "<p style='color:#888;'>Wczytaj plik XML, aby zobaczyć podgląd faktury.</p></body>")

    def extract_ksef_from_filename(self, path):
        filename = os.path.basename(path)
        match = re.search(r'(\d{10}-\d{8}-[A-Z0-9]{12,16}-\d{2})', filename)
        if match: return match.group(1)
        match_alt = re.search(r'(\d{10}-\d{8}-[A-Z0-9-]+)', filename)
        return match_alt.group(1) if match_alt else "Brak numeru KSeF"

    def get_val(self, root, paths):
        for path in paths:
            el = root.find(path, namespaces=self.ns)
            if el is not None and el.text: return el.text.strip()
        return ""

    def load_xml(self):
        path = filedialog.askopenfilename(filetypes=[("XML KSeF", "*.xml")])
        if not path: return
        self.filename_ksef_id = self.extract_ksef_from_filename(path)
        try:
            tree = etree.parse(path)
            root = tree.getroot()
            ns_url = root.nsmap.get(None)
            self.ns = {"n": ns_url} if ns_url else {}

            self.invoice_data = {
                "numer": self.get_val(root, [".//n:Fa/n:P_2"]),
                "data_wyst": self.get_val(root, [".//n:Fa/n:P_1"]),
                "data_sprz": self.get_val(root, [".//n:Fa/n:P_6"]),
                "okres_do": self.get_val(root, [".//n:Fa/n:OkresFa/n:P_6_Do"]),
                "brutto": self.get_val(root, [".//n:Fa/n:P_15"]),
                "waluta": self.get_val(root, [".//n:Fa/n:KodWaluty"]) or "PLN",
                "s_nazwa": self.get_val(root, [".//n:Podmiot1/n:DaneIdentyfikacyjne/n:Nazwa"]),
                "s_nip": self.get_val(root, [".//n:Podmiot1/n:DaneIdentyfikacyjne/n:NIP"]),
                "s_prefix": self.get_val(root, [".//n:Podmiot1/n:DaneIdentyfikacyjne/n:KodKrajuWydaniaNIP"]) or "PL",
                "s_adres": f"{self.get_val(root, ['.//n:Podmiot1/n:Adres/n:AdresL1'])}",
                "s_email": self.get_val(root, [".//n:Podmiot1/n:DaneKontaktowe/n:Email"]),
                "s_tel": self.get_val(root, [".//n:Podmiot1/n:DaneKontaktowe/n:Telefon"]),
                "n_nazwa": self.get_val(root, [".//n:Podmiot2/n:DaneIdentyfikacyjne/n:Nazwa"]),
                "n_nip": self.get_val(root, [".//n:Podmiot2/n:DaneIdentyfikacyjne/n:NIP"]),
                "n_adres": f"{self.get_val(root, ['.//n:Podmiot2/n:Adres/n:AdresL1'])}",
                "bank_nr": self.get_val(root, [".//n:Platnosc/n:RachunekBankowy/n:NrRB"]),
                "bank_nazwa": self.get_val(root, [".//n:Platnosc/n:RachunekBankowy/n:NazwaBanku"]),
                "termin": self.get_val(root, [".//n:Platnosc/n:TerminPlatnosci/n:Termin"]),
                "forma_platnosci": "Przelew",
                "regon": self.get_val(root, [".//n:Stopka/n:Rejestry/n:REGON"]),
            }

            self.podmioty3 = []
            for p3 in root.findall(".//n:Podmiot3", namespaces=self.ns):
                rola_map = {"1": "Wystawca", "2": "Podmiot inny 1", "3": "Odbiorca", "7": "Podmiot inny 2"}
                rola_code = self.get_val(p3, ["n:Rola"])
                self.podmioty3.append({
                    "nazwa": self.get_val(p3, ["n:DaneIdentyfikacyjne/n:Nazwa"]),
                    "nip": self.get_val(p3, ["n:DaneIdentyfikacyjne/n:NIP"]) or self.get_val(p3, ["n:NrEORI"]) or "-",
                    "rola": rola_map.get(rola_code, f"Podmiot inny {rola_code}"),
                    "adres": f"{self.get_val(p3, ['n:Adres/n:AdresL1'])}"
                })

            self.invoice_items = []
            for wiersz in root.findall(".//n:FaWiersz", namespaces=self.ns):
                self.invoice_items.append({
                    "lp": self.get_val(wiersz, ["n:NrWierszaFa"]),
                    "nazwa": self.get_val(wiersz, ["n:P_7"]),
                    "cena_jedn": self.get_val(wiersz, ["n:P_9B"]) or self.get_val(wiersz, ["n:P_9A"]),
                    "ilosc": self.get_val(wiersz, ["n:P_8B"]),
                    "miara": self.get_val(wiersz, ["n:P_8A"]),
                    "stawka": self.get_val(wiersz, ["n:P_12"]),
                    "netto": self.get_val(wiersz, ["n:P_11A"]),
                    "indeks": self.get_val(wiersz, ["n:Indeks"]) or "-"
                })

            self.vat_summary = []
            for i in range(1, 10):
                v_netto = self.get_val(root, [f".//n:Fa/n:P_13_{i}"])
                v_vat = self.get_val(root, [f".//n:Fa/n:P_14_{i}"])
                if v_netto:
                    stawka_opis = "23% lub 22%" if i == 1 else (f"{i}%" if i != 2 else "8% lub 7%")
                    brutto = float(v_netto) + float(v_vat or 0)
                    self.vat_summary.append({
                        "lp": len(self.vat_summary) + 1,
                        "stawka": stawka_opis,
                        "netto": v_netto,
                        "vat": v_vat or "0.00",
                        "brutto": f"{brutto:.2f}"
                    })

            self.status_label.configure(text=f"Wczytano: {self.invoice_data['numer']}", text_color="#10b981")
            self.refresh_preview()
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie udało się przetworzyć XML: {e}")

    def refresh_preview(self):
        rows = "".join([
            f"<tr><td>{it['lp']}</td><td>{it['nazwa']}</td><td class='r'>{it['cena_jedn']}</td><td>{it['ilosc']}</td><td>{it['miara']}</td><td class='c'>{it['stawka']}</td><td class='r'>{it['netto']}</td><td>{it['indeks']}</td></tr>"
            for it in self.invoice_items])

        vat_rows = "".join([
            f"<tr><td>{v['lp']}</td><td>{v['stawka']}</td><td class='r'>{v['netto']}</td><td class='r'>{v['vat']}</td><td class='r'>{v['brutto']}</td></tr>"
            for v in self.vat_summary])

        p_inny_html = ""
        for p in self.podmioty3:
            p_inny_html += f"""
            <div class="section-header">{p['rola']}</div>
            <div class="grid">
                <div>NIP: {p['nip']}<br>Nazwa: {p['nazwa']}<br>Rola: {p['rola']}</div>
                <div>Adres<br>{p['adres']}</div>
            </div><hr>"""

        html = f"""
        <html>
        <head>
        <style>
            body {{ font-family: 'Arial', sans-serif; background: #f5f5f5; padding: 0; margin: 0; font-size: 11px; color: #111; }}
            .page {{ max-width: 900px; margin: 20px auto; background: white; padding: 50px; box-shadow: 0 0 10px rgba(0,0,0,0.1); position: relative; }}
            .header-top {{ display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 20px; }}
            .ksef-logo {{ font-size: 26px; font-weight: normal; color: #000; }}
            .ksef-logo span {{ color: #e30613; font-weight: bold; }}
            .top-info {{ text-align: right; }}
            .inv-number {{ font-size: 22px; font-weight: bold; margin-bottom: 5px; }}
            .ksef-id {{ font-size: 10px; color: #333; }}
            hr {{ border: 0; border-top: 1px solid #ddd; margin: 20px 0; }}
            .grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 40px; margin-bottom: 5px; }}
            .section-header {{ font-weight: bold; font-size: 14px; margin-bottom: 10px; margin-top: 5px; }}
            table {{ width: 100%; border-collapse: collapse; margin-bottom: 15px; border: 1px solid #ccc; }}
            th {{ background: #fcfcfc; text-align: left; padding: 8px; border: 1px solid #ccc; font-size: 10px; font-weight: bold; }}
            td {{ padding: 8px; border: 1px solid #ccc; vertical-align: top; }}
            .r {{ text-align: right; }} .c {{ text-align: center; }}
            .total-box {{ font-weight: bold; font-size: 15px; text-align: right; margin: 20px 0; }}
            .qr-section {{ margin-top: 50px; border-top: 1px solid #eee; padding-top: 20px; display: flex; gap: 30px; }}
            .qr-box {{ width: 140px; height: 140px; border: 1px solid #000; padding: 5px; display: flex; align-items: center; justify-content: center; font-size: 8px; background: #000; color: white; }}
        </style>
        </head>
        <body>
            <div class="page">
                <div class="header-top">
                    <div class="ksef-logo">Krajowy System <span>e</span>-Faktur</div>
                    <div class="top-info">
                        <div style="font-size: 10px;">Numer Faktury:</div>
                        <div class="inv-number">{self.invoice_data['numer']}</div>
                        <div style="font-size: 11px; margin-bottom: 5px;">Faktura podstawowa</div>
                        <div class="ksef-id">Numer KSEF: {self.filename_ksef_id}</div>
                    </div>
                </div>
                <hr>
                <div class="grid">
                    <div>
                        <div class="section-header">Sprzedawca</div>
                        Prefix VAT: {self.invoice_data['s_prefix']}<br>
                        NIP: {self.invoice_data['s_nip']}<br>
                        Nazwa: {self.invoice_data['s_nazwa']}<br><br>
                        <strong>Adres</strong><br>{self.invoice_data['s_adres']}<br>Polska<br><br>
                        Dane kontaktowe<br>
                        E-mail: {self.invoice_data['s_email']}<br>
                        Tel.: {self.invoice_data['s_tel']}
                    </div>
                    <div>
                        <div class="section-header">Nabywca</div>
                        NIP: {self.invoice_data['n_nip']}<br>
                        Nazwa: {self.invoice_data['n_nazwa']}<br><br>
                        <strong>Adres</strong><br>{self.invoice_data['n_adres']}<br>Polska
                    </div>
                </div>
                <hr>
                {p_inny_html}
                <div class="section-header">Szczegóły</div>
                <div class="grid">
                    <div>Data wystawienia, z zastrzeżeniem art. 106na ust. 1 ustawy: {self.invoice_data['data_wyst']}</div>
                    <div>Data dokonania lub zakończenia dostawy towarów lub wykonania usługi: {self.invoice_data['data_sprz'] or self.invoice_data['okres_do'] or '-'}</div>
                </div>
                <hr>
                <div class="section-header">Pozycje</div>
                <div style="margin-bottom: 5px; font-style: italic;">Faktura wystawiona w cenach netto w walucie {self.invoice_data['waluta']}</div>
                <table>
                    <thead><tr><th>Lp.</th><th>Nazwa towaru lub usługi</th><th>Cena jedn. netto</th><th>Ilość</th><th>Miara</th><th>Stawka podatku</th><th>Wartość sprzedaży netto</th><th>Indeks</th></tr></thead>
                    <tbody>{rows}</tbody>
                </table>
                <div class="total-box">Kwota należności ogółem: {self.invoice_data['brutto']} {self.invoice_data['waluta']}</div>
                <div class="section-header">Podsumowanie stawek podatku</div>
                <table>
                    <thead><tr><th>Lp.</th><th>Stawka podatku</th><th>Kwota netto</th><th>Kwota podatku</th><th>Kwota brutto</th></tr></thead>
                    <tbody>{vat_rows}</tbody>
                </table>
                <div class="section-header">Płatność</div>
                Informacja o płatności: Brak zapłaty<br>
                Forma płatności: {self.invoice_data['forma_platnosci']}
                <div style="margin-top: 10px; width: 300px; margin-left: auto;">
                    <table><thead><tr><th>Termin płatności</th></tr></thead><tbody><tr><td>{self.invoice_data['termin']}</td></tr></tbody></table>
                </div>
                <div class="section-header">Numer rachunku bankowego</div>
                <table style="width: 450px;">
                    <tr><td style="width: 150px; background: #f9f9f9;">Pełny numer rachunku</td><td>{self.invoice_data['bank_nr']}</td></tr>
                    <tr><td style="background: #f9f9f9;">Kod SWIFT</td><td>-</td></tr>
                    <tr><td style="background: #f9f9f9;">Rachunek własny banku</td><td>-</td></tr>
                    <tr><td style="background: #f9f9f9;">Nazwa banku</td><td>{self.invoice_data['bank_nazwa']}</td></tr>
                    <tr><td style="background: #f9f9f9;">Opis rachunku</td><td>-</td></tr>
                </table>
                <div class="section-header">Rejestry</div>
                <table style="width: 450px;">
                    <tr><td style="width: 150px; background: #f9f9f9;">REGON</td><td>{self.invoice_data['regon'] or '-'}</td></tr>
                </table>
                <div class="qr-section">
                   <!-- <div>
                        <div style="font-weight: bold; margin-bottom: 10px;">Sprawdź, czy Twoja faktura znajduje się w KSeF!</div>
                        <div class="qr-box">[KOD QR]</div>
                        <div style="font-size: 9px; margin-top: 10px; font-family: monospace;">{self.filename_ksef_id}</div>
                    </div>
                    <div style="flex: 1; font-size: 10px; color: #555; padding-top: 30px;">
                        Nie możesz zeskanować kodu z obrazka? Kliknij w link weryfikacyjny i przejdź do weryfikacji faktury!<br>
                        <a href="#" style="color: blue;">https://qr.ksef.mf.gov.pl/invoice/{self.filename_ksef_id}</a>
                    </div>-->
                </div>
            </div>
        </body>
        </html>
        """
        self.preview_frame.load_html(html)

    def set_cell_background(self, cell, fill_color):
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def export_word(self):
        if not self.invoice_data: return
        path = filedialog.asksaveasfilename(defaultextension=".docx",
                                            initialfile=f"Opis_{self.invoice_data['numer'].replace('/', '_')}.docx")
        if not path: return

        doc = Document()
        title = doc.add_paragraph()
        run = title.add_run("Opis merytoryczny faktury KSeF")
        run.bold = True
        run.font.size = Pt(14)

        def add_section_header(text):
            table = doc.add_table(rows=1, cols=1)
            table.width = Inches(6.5)
            cell = table.rows[0].cells[0]
            cell.text = text
            self.set_cell_background(cell, "003399")
            run = cell.paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
            run.font.size = Pt(10)

        add_section_header("CZ. I. PRACOWNIK MERYTORYCZNY")
        t1 = doc.add_table(rows=11, cols=2)
        t1.style = 'Table Grid'

        def fill_row(idx, l1, v1, l2, v2):
            t1.rows[idx].cells[0].text = f"{l1}\n{v1}"
            t1.rows[idx].cells[1].text = f"{l2}\n{v2}"

        fill_row(0, "Komórka odpowiedzialna:", "[...................]", "Data wpływu faktury:",
                 self.invoice_data['data_wyst'])
        fill_row(1, "Dokument na podstawie:", "[ Zapotrzebowanie ]", "Nr umowy / wniosku:", "...................")
        fill_row(2, "* Numer KSeF:", self.filename_ksef_id, "* Numer faktury:", self.invoice_data['numer'])
        fill_row(3, "* Data wystawienia:", self.invoice_data['data_wyst'], "Termin płatności:",
                 self.invoice_data['termin'])
        fill_row(4, "* Sprzedawca:", self.invoice_data['s_nazwa'], "* NIP:", self.invoice_data['s_nip'])
        addr_c = t1.rows[5].cells[0];
        addr_c.merge(t1.rows[5].cells[1]);
        addr_c.text = f"Adres:\n{self.invoice_data['s_adres']}"
        fill_row(6, "Rachunek bankowy:", self.invoice_data['bank_nr'], "* Kwota brutto:",
                 f"{self.invoice_data['brutto']} {self.invoice_data['waluta']}")

        vat_info = "".join([f"{v['stawka']}: {v['brutto']} " for v in self.vat_summary])
        vat_c = t1.rows[7].cells[0];
        vat_c.merge(t1.rows[7].cells[1]);
        vat_c.text = f"Podsumowanie VAT:\n{vat_info}"
        desc_c = t1.rows[8].cells[0];
        desc_c.merge(t1.rows[8].cells[1]);
        desc_c.text = "* Opis wydatku:\n\n\n"
        fill_row(9, "Protokół odbioru:", "[ Tak / Nie ]", "Uwagi:", "...................")
        fill_row(10, "Zrealizowano zgodnie z umową:", "[ Tak / Nie ]", "Uwagi:", "...................")

        doc.add_paragraph()
        add_section_header("CZ. II. ZAMÓWIENIA PUBLICZNE")
        t2 = doc.add_table(rows=2, cols=2);
        t2.style = 'Table Grid'
        t2.rows[0].cells[0].text = "Wydatek bez PZP:\n[ ] Tak [ ] Nie"
        t2.rows[1].cells[0].merge(t2.rows[1].cells[1]);
        t2.rows[1].cells[0].text = "Podpis: ........................"

        doc.add_paragraph()
        add_section_header("CZ. III. KSIĘGOWOŚĆ")
        t3 = doc.add_table(rows=2, cols=4);
        t3.style = 'Table Grid'
        h = t3.rows[0].cells;
        h[0].text, h[1].text, h[2].text, h[3].text = "Dział", "Rozdział", "Paragraf", "Kwota"
        t3.rows[1].cells[3].text = self.invoice_data['brutto']

        doc.add_paragraph()
        add_section_header("CZ. IV. ZATWIERDZENIE")
        doc.add_paragraph("\nPodpis Skarbnika: ....................   Podpis Kierownika: ....................")

        doc.save(path)
        messagebox.showinfo("Sukces", "Opis Word wygenerowany pomyślnie.")


if __name__ == "__main__":
    app = KSeFApp()
    app.mainloop()
